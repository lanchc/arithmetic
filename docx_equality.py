# -*- coding:utf-8 _*-
# @time     2021/12/3 5:03 下午
# @explain  等式成立
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from dir_helper import *
from math_helper import *

# 构建特殊题目
#  3 +  7 =  9 + （   ） 

# 计算结果
def compute_outcome_back(a, b, operator):
    if operator == '-':
        return a + b
    else:
        return a - b


# 格式化输出模版
def format_basics_text(a, b, c, d, sym_a, sym_b):
    style_rd = random.randint(1, 4)
    style_kh = '（   ）'
    if 1 == style_rd:
        return '{0: >2} {1:} {2: >2} = {3: >2} {4:} {5: >2} '.format(int(a), str(sym_a), int(b), str(style_kh), str(sym_b), int(d))
    elif 2 == style_rd:
        return '{0: >2} {1:} {2: >2} = {3: >2} {4:} {5: >2} '.format(int(a), str(sym_a), str(style_kh), int(c), str(sym_b), int(d))
    elif 3 == style_rd:
        return '{0: >2} {1:} {2: >2} = {3: >2} {4:} {5: >2} '.format(int(a), str(sym_a), int(b), int(c), str(sym_b), str(style_kh))
    elif 4 == style_rd:
        return '{0: >2} {1:} {2: >2} = {3: >2} {4:} {5: >2} '.format(str(style_kh), str(sym_a), int(b), int(c), str(sym_b), int(d))
    else:
        return '{0: >2} {1:} {2: >2} = {3: >2} {4:} {5: >2} '.format(int(a), str(sym_a), int(b), int(c), str(sym_b), int(d))


# range_number 单个数字上限
# max_number 计算和上限
def random_subject(range_number=10, max_number=10):

    is_complete = True
    topic_str = ''

    while is_complete:
        is_skip = True

        symbol_array = ['-', '+']
        temp_l_sym = random.choice(symbol_array)
        temp_r_sym = random.choice(symbol_array)
        # x =（）- 3
        tmp_aa = random.randint(0, range_number)
        tmp_bb = random.randint(0, range_number)
        tmp_cc = random.randint(0, range_number)
        # 左侧结果
        temp_zz = compute_outcome(tmp_aa, tmp_bb, temp_l_sym)
        # 朝纲过滤 - 被减数小于减数
        if temp_l_sym == '-' and tmp_aa < tmp_bb:
            is_skip = False

        # 朝纲啦 限制最大值
        if temp_zz > max_number:
            is_skip = False

        # 等式无法构建
        if temp_r_sym == '+' and tmp_cc > temp_zz:
            is_skip = False

        # 等式无法构建
        if temp_r_sym == '-' and tmp_cc < temp_zz:
            is_skip = False

        # 计算最后一个数据
        tmp_dd = compute_outcome_back(temp_zz, tmp_cc, temp_r_sym)

        # 朝纲啦 限制最大值
        if tmp_dd > max_number:
            is_skip = False

        # 构建题目
        if is_skip:
            if tmp_cc > tmp_dd:
                topic_str = format_basics_text(tmp_aa, tmp_bb, tmp_cc, tmp_dd, temp_l_sym, temp_r_sym)
            else:
                topic_str = format_basics_text(tmp_aa, tmp_bb, tmp_dd, tmp_cc, temp_l_sym, temp_r_sym)
            is_complete = False

    return topic_str


# 构建数据
def generate_topic_mingle(number=75, row_max=3, template_name='template_more.docx'):
    # 读取文档
    template_docx = Document(template_path(template_name))
    # 字体样式
    template_docx.styles['Normal'].font.name = u'Courier New'
    template_docx.styles['Normal'].font.size = Pt(10)
    template_docx.styles['Normal'].font.color.rgb = RGBColor(33, 33, 33)
    # 首个表格
    write_table = template_docx.tables[0]
    # 跳过标题
    write_next_row = 1
    write_next_tag = row_max - 1

    # 插入更新内容
    for k in range(number):
        remainder = k % row_max
        # print("remainder=", remainder)
        r_cells = write_table.rows[write_next_row].cells
        r_cells[remainder].text = random_subject()
        r_cells[remainder].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 四题下移
        if write_next_tag == remainder:
            write_next_row += 1

    # 保存文件
    template_docx.save(docx_rename())


if __name__ == '__main__':

    # 检查导出文件夹是否存在
    checkup_dir(project_root_dir() + '/dist/')

    # print(random_subject())
    # print(random_subject())
    # print(random_subject())
    # print(random_subject())


    # 构建文件份数
    documents_count = 20

    print("文档操作助手")
    for i in range(1, documents_count + 1):
        print("已经完成: " + str(i) + "/" + str(documents_count))
        generate_topic_mingle()
        time.sleep(1)
    print("创建操作完成")

