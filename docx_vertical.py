# -*- coding:utf-8 _*-
# @time     2021/12/3 5:03 下午
# @explain  竖式计算
from docx import Document
from docx.shared import Pt, RGBColor
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from dir_helper import *
from math_helper import *
import time


# range_number 单个数字上限
# max_number 计算和上限
def random_subject(range_number=100, max_number=100):

    is_complete = True
    topic_str = ''

    while is_complete:
        is_skip = True

        symbol_array = ['-', '+']
        temp_l_sym = random.choice(symbol_array)
        # x =（）- 3
        tmp_aa = random.randint(10, range_number)
        tmp_bb = random.randint(10, range_number)
        # 左侧结果
        temp_zz = compute_outcome(tmp_aa, tmp_bb, temp_l_sym)

        # 相同数据移除
        if tmp_aa == tmp_bb:
            is_skip = False

        # 朝纲过滤 - 被减数小于减数
        if temp_l_sym == '-' and tmp_aa < tmp_bb:
            is_skip = False

        # 朝纲啦 限制最大值
        if temp_zz > max_number:
            is_skip = False


        # 构建题目
        if is_skip:
            topic_str = '{0: >2} {1:} {2: >2} =  '.format(int(tmp_aa), str(temp_l_sym), int(tmp_bb))
            is_complete = False

    return topic_str

# 构建数据
def generate_topic_mingle(number=20, row_max=4, template_name='vertically.docx'):
    # 读取文档
    template_docx = Document(template_path(template_name))
    # 字体样式
    template_docx.styles['Normal'].font.name = u'Courier New'
    template_docx.styles['Normal'].font.size = Pt(10)
    template_docx.styles['Normal'].font.color.rgb = RGBColor(33, 33, 33)
    # 首个表格
    write_table = template_docx.tables[0]
    # 跳过标题
    write_next_row = 2
    write_next_tag = row_max - 1

    # 插入更新内容
    for k in range(number):
        remainder = k % row_max
        r_cells = write_table.rows[write_next_row].cells
        r_cells[remainder].text = random_subject()
        r_cells[remainder].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT

        # 四题下移
        if write_next_tag == remainder:
            write_next_row += 2

    # 保存文件
    template_docx.save(docx_rename())


if __name__ == '__main__':
    pass
    # 检查导出文件夹是否存在
    checkup_dir(project_root_dir() + '/dist/')

    # 构建文件份数
    documents_count = 30

    print("文档操作助手")
    for i in range(1, documents_count + 1):
        print("已经完成: " + str(i) + "/" + str(documents_count))
        generate_topic_mingle()
        time.sleep(1)
    print("创建操作完成")
