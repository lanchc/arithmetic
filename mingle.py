# -*- coding:utf-8 _*-
# @time     2021/11/10 12:41 下午
# @explain  arithmetic
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from dir_helper import *
from math_helper import *


# 格式化输出模版 - 混合运算
def format_mingle_text(a, b, c, d, sy_a, sy_b):
    # random_blank = random.randint(1, 4)
    random_blank = 4
    blanks_blank = "（   ）"
    if 1 == random_blank:
        return '{0:} {1:} {2: >2} {3:} {4: >2} = {5: >2}' \
            .format(str(blanks_blank), str(sy_a), int(b), str(sy_b), int(c), int(d))
    elif 2 == random_blank:
        return '{0: >2} {1:} {2: >2} {3:} {4: >2} = {5: >2}' \
            .format(int(a), str(sy_a), str(blanks_blank), str(sy_b), int(c), int(d))
    elif 3 == random_blank:
        return '{0: >2} {1:} {2: >2} {3:} {4:} = {5: >2}' \
            .format(int(a), str(sy_a), int(b), str(sy_b), str(blanks_blank), int(d))
    elif 4 == random_blank:
        return '{0: >2} {1:} {2: >2} {3:} {4: >2} = {5:}' \
            .format(int(a), str(sy_a), int(b), str(sy_b), int(c), str(blanks_blank))
    else:
        return '{0: >2} {1:} {2: >2} {3:} {4: >2} = {5: >2}' \
            .format(int(a), str(sy_a), int(b), str(sy_b), int(c), int(d))


# 随机题目
def random_mingle_topic(answer_max=10, limit_upper=20):

    is_complete = True
    topic_str = ''

    while is_complete:
        is_skip = True

        symbol_array = ['-', '+']
        # 随机符号
        kw_sy_a = random.choice(symbol_array)
        # 随机符号
        kw_sy_b = random.choice(symbol_array)
        # 随机参数
        kw_aa = random.randint(0, answer_max)
        kw_bb = random.randint(0, answer_max)
        kw_cc = random.randint(0, answer_max)

        kw_dd = compute_outcome_mingle(kw_aa, kw_bb, kw_cc, kw_sy_a, kw_sy_b)

        # 朝纲过滤 - 被减数小于减数
        if kw_sy_a == '-' and kw_aa < kw_bb:
            is_skip = False

        # 朝纲过滤 - 总和超过限制
        if kw_dd > limit_upper:
            is_skip = False

        # 朝纲过滤 - 负数
        if kw_dd < 0:
            is_skip = False

        # 构建题目
        if is_skip:
            topic_str = format_mingle_text(kw_aa, kw_bb, kw_cc, kw_dd, kw_sy_a, kw_sy_b)
            is_complete = False

    return topic_str


# 构建数据
def generate_topic_mingle(number=69, row_max=3, template_name='template_more.docx'):
    # 读取文档
    template_docx = Document(template_path(template_name))
    # 字体样式
    # template_docx.styles['Normal'].font.name = u'Times New Roman'
    # template_docx.styles['Normal'].font.size = Pt(10.5)
    # template_docx.styles['Normal'].font.color.rgb = RGBColor(33, 33, 33)
    # 首个表格
    write_table = template_docx.tables[0]
    # 跳过标题
    write_next_row = 1
    write_next_tag = row_max - 1

    # 插入更新内容
    for k in range(number):
        remainder = k % row_max
        # print("remainder=", remainder)
        r_cells = write_table.rows[write_next_row].cells
        r_cells[remainder].text = random_mingle_topic()
        r_cells[remainder].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 四题下移
        if write_next_tag == remainder:
            write_next_row += 1

    # 保存文件
    template_docx.save(docx_rename())


if __name__ == '__main__':

    # 检查导出文件夹是否存在
    checkup_dir(project_root_dir() + '/dist/')

    # 构建文件份数
    documents_count = 50

    print("文档操作助手")
    for i in range(1, documents_count + 1):
        print("已经完成: " + str(i) + "/" + str(documents_count))
        generate_topic_mingle()
        time.sleep(1)
    print("创建操作完成")

