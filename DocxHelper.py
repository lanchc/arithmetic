# -*- coding:utf-8 _*-
# @time     2021/10/29 9:02 上午
# @explain  arithmetic
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from dir_helper import *
from math_helper import *


# 随机题目
def random_topic(answer_max=20):

    is_complete = True
    topic_str = ''

    while is_complete:
        is_skip = True

        symbol_array = ['-', '+']
        kw_sy = random.choice(symbol_array)
        kw_aa = random.randint(0, answer_max)
        kw_bb = random.randint(0, answer_max)
        kw_cc = compute_outcome(kw_aa, kw_bb, kw_sy)
        # 朝纲过滤 - 被减数小于减数
        if kw_sy == '-' and kw_aa < kw_bb:
            is_skip = False
        # 朝纲过滤 - 总和超过限制
        if kw_cc > answer_max:
            is_skip = False

        # 构建题目
        if is_skip:
            topic_str = format_basics_text(kw_aa, kw_sy, kw_bb, kw_cc)
            is_complete = False

    return topic_str


# 构建数据
def generate_topic(number=100, answer_max=20):
    # 读取文档
    template_docx = Document(template_path())
    # 字体样式
    template_docx.styles['Normal'].font.name = u'Courier New'
    template_docx.styles['Normal'].font.size = Pt(10)
    template_docx.styles['Normal'].font.color.rgb = RGBColor(33, 33, 33)
    # 首个表格
    write_table = template_docx.tables[0]
    # 跳过标题
    write_next_row = 2

    # 插入更新内容
    for k in range(number):
        remainder = k % 4

        r_cells = write_table.rows[write_next_row].cells
        r_cells[remainder].text = random_topic(answer_max)
        r_cells[remainder].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 四题下移
        if 3 == remainder:
            write_next_row += 1

    # 保存文件
    template_docx.save(docx_rename())


if __name__ == '__main__':

    # 检查导出文件夹是否存在
    checkup_dir(project_root_dir() + '/dist/')

    # 构建文件份数
    documents_count = 50

    print("文档操作助手")
    for i in range(1, documents_count + 1):
        print("已经完成: " + str(i) + "/" + str(documents_count))
        generate_topic()
        time.sleep(1)
    print("创建操作完成")
